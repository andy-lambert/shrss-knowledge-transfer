#!/usr/bin/env python3
"""
Consolidate all session transcript .docx files in this directory into a single
Word document, with a page break at the end of each session's transcript.
Excludes lock files (~$*) and the output file. Strips images and image references.
"""
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# OOXML tags for image content: remove these from copied body elements
W_DRAWING = qn("w:drawing")
W_PICT = qn("w:pict")


def strip_images_from_element(element) -> None:
    """Remove w:drawing and w:pict descendants from element (in-place)."""
    to_remove = [
        child
        for child in element.iter()
        if child.tag in (W_DRAWING, W_PICT)
    ]
    for el in to_remove:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def add_page_break(doc: Document) -> None:
    """Append a page break paragraph to the document."""
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._r.append(br)


def main() -> None:
    transcript_dir = Path(__file__).resolve().parent
    out_name = "SHRSS_Adobe_KT_All_Session_Transcripts_Consolidated.docx"
    out_path = transcript_dir / out_name

    docx_files = sorted(
        f
        for f in transcript_dir.glob("*.docx")
        if not f.name.startswith("~$") and f.name != out_name
    )

    if not docx_files:
        print("No .docx files found to consolidate.")
        return

    merged = Document()
    # Optional: add a title paragraph
    merged.add_paragraph(
        "SHRSS Adobe Knowledge Transfer — All Session Transcripts (Consolidated)",
        style="Title",
    )
    merged.add_paragraph()

    for i, path in enumerate(docx_files):
        print(f"Adding: {path.name}")
        try:
            sub = Document(path)
        except Exception as e:
            print(f"  Skip ({e})")
            continue
        # Copy each body element (paragraph, table, etc.) into the merged doc; strip images
        for element in list(sub.element.body):
            copied = deepcopy(element)
            strip_images_from_element(copied)
            merged.element.body.append(copied)
        # Page break at the end of each session (including after the last, for consistency)
        add_page_break(merged)

    merged.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    main()
